"""
Handles turning raw uploads (PDF/DOCX/TXT/MD/CSV) or a website URL into
chunked, embedded documents stored in a bot's vector collection.
"""
import os
from typing import List

import requests
from bs4 import BeautifulSoup
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

from app.config import settings
from app.database import vector_store_manager
from app.utils.exceptions import UnsupportedFileTypeError, FileTooLargeError, IngestionError
from app.utils.logger import get_logger

logger = get_logger(__name__)

_splitter = RecursiveCharacterTextSplitter(
    chunk_size=settings.CHUNK_SIZE,
    chunk_overlap=settings.CHUNK_OVERLAP,
    separators=["\n\n", "\n", ". ", " ", ""],
)


def _load_text_from_file(file_path: str, ext: str) -> str:
    try:
        if ext == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            return "\n".join((page.extract_text() or "") for page in reader.pages)
        elif ext == ".docx":
            import docx
            doc = docx.Document(file_path)
            return "\n".join(p.text for p in doc.paragraphs)
        elif ext in (".txt", ".md", ".csv"):
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        else:
            raise UnsupportedFileTypeError(ext)
    except UnsupportedFileTypeError:
        raise
    except Exception as e:
        logger.error(f"Failed to parse file {file_path}: {e}")
        raise IngestionError(f"could not parse {os.path.basename(file_path)} ({e})")


def validate_file(filename: str, size_bytes: int):
    ext = os.path.splitext(filename)[1].lower()
    if ext not in settings.ALLOWED_EXTENSIONS:
        raise UnsupportedFileTypeError(ext)
    max_bytes = settings.MAX_FILE_SIZE_MB * 1024 * 1024
    if size_bytes > max_bytes:
        raise FileTooLargeError(settings.MAX_FILE_SIZE_MB)
    return ext


def ingest_file(bot_id: str, file_path: str, original_filename: str) -> int:
    ext = os.path.splitext(original_filename)[1].lower()
    raw_text = _load_text_from_file(file_path, ext)
    if not raw_text or not raw_text.strip():
        raise IngestionError("no extractable text found in the document")

    chunks = _splitter.split_text(raw_text)
    documents = [
        Document(page_content=chunk, metadata={"source": original_filename, "bot_id": bot_id})
        for chunk in chunks
    ]
    added = vector_store_manager.add_documents(bot_id, documents)
    logger.info(f"Ingested '{original_filename}' for bot {bot_id}: {added} chunks")
    return added


def ingest_url(bot_id: str, url: str) -> int:
    try:
        resp = requests.get(str(url), timeout=15, headers={"User-Agent": "Mozilla/5.0 (RAG-Chatbot-Ingestor)"})
        resp.raise_for_status()
    except requests.RequestException as e:
        raise IngestionError(f"could not fetch URL ({e})")

    try:
        soup = BeautifulSoup(resp.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
        text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    except Exception as e:
        raise IngestionError(f"could not parse page content ({e})")

    if not text.strip():
        raise IngestionError("no extractable text found on the page")

    chunks = _splitter.split_text(text)
    documents = [
        Document(page_content=chunk, metadata={"source": str(url), "bot_id": bot_id})
        for chunk in chunks
    ]
    added = vector_store_manager.add_documents(bot_id, documents)
    logger.info(f"Ingested URL '{url}' for bot {bot_id}: {added} chunks")
    return added
